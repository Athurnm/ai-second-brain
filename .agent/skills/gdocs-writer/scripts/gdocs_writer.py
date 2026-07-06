#!/usr/bin/env python3
"""
Google Docs Writer Skill
Converts markdown to a clean .docx and uploads to Google Drive.

Rules:
- No em-dashes (—) — replaced with commas
- No double-dashes (--) — replaced with commas
- Numbered lists always on separate lines
- Bullet lists always on separate lines
"""

import os, sys, re, argparse, tempfile

SKILL_DIR        = os.path.dirname(os.path.abspath(__file__))
DEFAULT_CRED_DIR = os.path.join(SKILL_DIR, '..', '..', 'work-drive-connector')
SCOPES           = ['https://www.googleapis.com/auth/drive']

def authenticate(cred_dir=None):
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    token_file = os.path.join(cred_dir or DEFAULT_CRED_DIR, 'token.json')
    creds = None
    if os.path.exists(token_file):
        creds = Credentials.from_authorized_user_file(token_file, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            print(f"Error: Cannot authenticate. Token not found or expired: {token_file}")
            sys.exit(1)
    return creds

# ─── Text sanitiser ───────────────────────────────────────────────────────────
def sanitise(text):
    """Remove em-dashes and double-dashes, replace with comma+space."""
    # Em-dash with spaces around it
    text = re.sub(r'\s*—\s*', ', ', text)
    # Double-dash with spaces (avoid hitting table separator rows like |---|)
    text = re.sub(r'(?<!\|)\s*--\s*(?!\|)', ', ', text)
    # Clean up double commas or leading/trailing comma artifacts
    text = re.sub(r',\s*,', ',', text)
    text = re.sub(r'^,\s*', '', text)
    return text

# ─── Inline formatter (bold, italic) ──────────────────────────────────────────
def add_formatted_run(para, text):
    """Parse **bold**, *italic*, and plain text into docx runs."""
    from docx.shared import Pt
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)')
    for match in pattern.finditer(text):
        chunk = match.group(0)
        if chunk.startswith('**') and chunk.endswith('**'):
            run = para.add_run(sanitise(chunk[2:-2]))
            run.bold = True
        elif chunk.startswith('*') and chunk.endswith('*'):
            run = para.add_run(sanitise(chunk[1:-1]))
            run.italic = True
        else:
            para.add_run(sanitise(chunk))

# ─── Table parser ─────────────────────────────────────────────────────────────
def is_table_row(line):
    return line.strip().startswith('|') and line.strip().endswith('|')

def is_separator_row(line):
    return bool(re.match(r'^\|[\s\-\|:]+\|$', line.strip()))

def parse_table_row(line):
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]

# ─── Markdown → docx ──────────────────────────────────────────────────────────
def md_to_docx(md_path, title=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for s_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        s = doc.styles[s_name]
        s.font.name = 'Calibri'

    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')
    i = 0

    def set_para_spacing(para, before=0, after=6):
        from docx.shared import Pt
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)

    while i < len(lines):
        line = lines[i]

        # ── Blank line ──
        if line.strip() == '':
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^-{3,}$', line.strip()):
            p = doc.add_paragraph()
            set_para_spacing(p, before=4, after=4)
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pb = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pb.append(bottom)
            pPr.append(pb)
            i += 1
            continue

        # ── Headings ──
        h_match = re.match(r'^(#{1,3})\s+(.*)', line)
        if h_match:
            level  = len(h_match.group(1))
            text   = sanitise(h_match.group(2).strip())
            # Strip bold markers from heading text
            text   = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            h_name = f'Heading {level}'
            p = doc.add_heading(text, level=level)
            set_para_spacing(p, before=16 if level == 1 else 12, after=6)
            i += 1
            continue

        # ── Table ──
        if is_table_row(line):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_separator_row(lines[i]):
                    table_lines.append(parse_table_row(lines[i]))
                i += 1
            if not table_lines:
                continue
            num_cols = len(table_lines[0])
            tbl = doc.add_table(rows=len(table_lines), cols=num_cols)
            tbl.style = 'Table Grid'
            for r_idx, row_data in enumerate(table_lines):
                row = tbl.rows[r_idx]
                for c_idx, cell_text in enumerate(row_data[:num_cols]):
                    cell = row.cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_formatted_run(p, sanitise(cell_text))
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
                        # Header row shading
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F0F0F0')
                        tcPr.append(shd)
            # Spacing after table
            p = doc.add_paragraph()
            set_para_spacing(p, before=0, after=4)
            continue

        # ── Blockquote ──
        if line.strip().startswith('>'):
            text = sanitise(line.strip().lstrip('>').strip())
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            set_para_spacing(p, before=4, after=4)
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # ── Numbered list ──
        if re.match(r'^\s*\d+\.\s', line):
            while i < len(lines) and re.match(r'^\s*\d+\.\s', lines[i]):
                text = sanitise(re.sub(r'^\s*\d+\.\s+', '', lines[i]))
                p = doc.add_paragraph(style='List Number')
                add_formatted_run(p, text)
                set_para_spacing(p, before=2, after=2)
                i += 1
            # Extra spacing after list block
            p = doc.add_paragraph()
            set_para_spacing(p, before=0, after=2)
            continue

        # ── Bullet list ──
        if re.match(r'^\s*[-*]\s', line):
            while i < len(lines) and re.match(r'^\s*[-*]\s', lines[i]):
                text = sanitise(re.sub(r'^\s*[-*]\s+', '', lines[i]))
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_run(p, text)
                set_para_spacing(p, before=2, after=2)
                i += 1
            p = doc.add_paragraph()
            set_para_spacing(p, before=0, after=2)
            continue

        # ── Regular paragraph ──
        text = sanitise(line)
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        set_para_spacing(p, before=0, after=8)
        add_formatted_run(p, text)
        i += 1

    # ── Save ──
    fd, out_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(out_path)
    return out_path

# ─── Drive upload / update ────────────────────────────────────────────────────
def drive_upload(docx_path, title, share=True, cred_dir=None, parent_id=None):
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    creds   = authenticate(cred_dir)
    service = build('drive', 'v3', credentials=creds)

    mime_docx  = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    mime_gdoc  = 'application/vnd.google-apps.document'
    metadata   = {'name': title, 'mimeType': mime_gdoc}
    if parent_id:
        metadata['parents'] = [parent_id]
    media      = MediaFileUpload(docx_path, mimetype=mime_docx)

    file = service.files().create(
        body=metadata, media_body=media, fields='id, webViewLink'
    ).execute()

    file_id = file.get('id')
    link    = file.get('webViewLink')

    if share:
        service.permissions().create(
            fileId=file_id,
            body={'type': 'anyone', 'role': 'commenter'},
            fields='id'
        ).execute()

    return file_id, link

def drive_update(file_id, docx_path, cred_dir=None):
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    creds   = authenticate(cred_dir)
    service = build('drive', 'v3', credentials=creds)

    mime_docx = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    media     = MediaFileUpload(docx_path, mimetype=mime_docx)

    file = service.files().update(
        fileId=file_id,
        media_body=media,
        fields='id, webViewLink'
    ).execute()

    return file.get('id'), file.get('webViewLink')

# ─── CLI ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description='Google Docs Writer: md → docx → Google Drive')
    sub = parser.add_subparsers(dest='command')

    up = sub.add_parser('upload', help='Upload new Google Doc')
    up.add_argument('--file',      required=True, help='Path to .md file')
    up.add_argument('--title',     help='Document title (default: filename)')
    up.add_argument('--no-share',  action='store_true', help='Skip public comment permission')
    up.add_argument('--cred-dir',  help='Path to folder containing token.json (default: work-drive-connector)')
    up.add_argument('--parent-id', help='Google Drive folder ID to upload into')

    ud = sub.add_parser('update', help='Update existing Google Doc')
    ud.add_argument('--file',     required=True, help='Path to .md file')
    ud.add_argument('--id',       required=True, help='Google Doc file ID to update')
    ud.add_argument('--cred-dir', help='Path to folder containing token.json')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    md_path = args.file
    if not os.path.exists(md_path):
        print(f"Error: file not found: {md_path}")
        sys.exit(1)

    # Derive title
    base  = os.path.splitext(os.path.basename(md_path))[0]
    title = getattr(args, 'title', None) or base.replace('_', ' ')

    print(f"[gdocs-writer] Converting {md_path} ...")
    docx_path = md_to_docx(md_path, title=title)

    cred_dir = getattr(args, 'cred_dir', None)

    if args.command == 'upload':
        share     = not getattr(args, 'no_share', False)
        parent_id = getattr(args, 'parent_id', None)
        file_id, link = drive_upload(docx_path, title=title, share=share,
                                     cred_dir=cred_dir, parent_id=parent_id)
        print(f"[gdocs-writer] Uploaded: {title}")
        print(f"File ID: {file_id}")
        print(f"Link: {link}")

    elif args.command == 'update':
        file_id, link = drive_update(args.id, docx_path, cred_dir=cred_dir)
        print(f"[gdocs-writer] Updated: {args.id}")
        print(f"File ID: {file_id}")
        print(f"Link: {link}")

    os.unlink(docx_path)

if __name__ == '__main__':
    main()
